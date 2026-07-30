"""파일 타입별 텍스트 추출 모듈."""

import io
import logging
import os
import re
import unicodedata
import xml.etree.ElementTree as ET
import zipfile
from contextlib import closing
from pathlib import Path

import fitz  # PyMuPDF: 스캔 PDF를 이미지로 렌더링해 OCR 하기 위함
import pdfplumber
import pytesseract
from dotenv import load_dotenv
from PIL import Image, ImageOps

# 이 모듈이 pipeline 의 load_dotenv() 보다 먼저 임포트되므로 여기서 직접 .env 를 읽는다
# (classify.py 와 같은 이유). 안 그러면 OCR_PROVIDER 등이 기본값으로 떨어진다.
load_dotenv()

logger = logging.getLogger("archive_pipeline")

PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}
TEXT_EXTENSIONS = {".txt", ".md"}
HWP_EXTENSIONS = {".hwp"}  # 한글 v5 바이너리(OLE)
HWPX_EXTENSIONS = {".hwpx"}  # 한글 OWPML(zip+xml)
XLSX_EXTENSIONS = {".xlsx", ".xls"}
DOCX_EXTENSIONS = {".docx"}
MSG_EXTENSIONS = {".msg"}
XML_EXTENSIONS = {".xml"}  # 전자세금계산서(NTS) 등 구조화 XML

# 파이프라인이 받아들이는 전체 확장자 집합(단일 진실 공급원).
# run_once.py·watch.py 가 _inbox 를 거를 때 이 집합을 임포트해 쓴다 — 여기에만 추가하면
# 두 진입점이 동시에 새 형식을 인식한다.
SUPPORTED_EXTENSIONS = (
    PDF_EXTENSIONS
    | IMAGE_EXTENSIONS
    | TEXT_EXTENSIONS
    | HWP_EXTENSIONS
    | HWPX_EXTENSIONS
    | XLSX_EXTENSIONS
    | DOCX_EXTENSIONS
    | MSG_EXTENSIONS
    | XML_EXTENSIONS
)

OCR_LANGUAGES = "kor+eng"

# OCR 백엔드: tesseract(로컬 설치 필요, 기본) 또는 gemini(클라우드, Tesseract 없는 저사양 PC용).
# .env 의 OCR_PROVIDER 로 머신별로 고른다 — STRX-D75 는 기본값 그대로 Tesseract 를 쓴다.
OCR_PROVIDER = os.getenv("OCR_PROVIDER", "tesseract").lower()
# 스캔 OCR 은 분류보다 비전 성능이 중요해 분류용(GEMINI_MODEL)과 따로 고른다.
GEMINI_OCR_MODEL = os.getenv("GEMINI_OCR_MODEL", "gemini-3.5-flash-lite")
# 페이지마다 API 를 1회 호출하므로, 무인 운영 중 수백 장짜리 스캔 하나가 요금을
# 폭주시키지 않도록 상한을 둔다. 넘으면 앞쪽 페이지까지만 읽고 경고를 남긴다.
GEMINI_OCR_MAX_PAGES = int(os.getenv("GEMINI_OCR_MAX_PAGES", "30"))
# 한 페이지 전사는 분류용 JSON 보다 훨씬 길다. 상한을 안 주면 API 기본값에서 조용히 잘리고,
# 잘린 줄이 '문서 전체'로 넘어간다. 넉넉히 잡되, 그래도 잘리면 아래에서 예외로 올린다.
OCR_MAX_OUTPUT_TOKENS = int(os.getenv("GEMINI_OCR_MAX_OUTPUT_TOKENS", "8192"))

# pyhwp 가 내부적으로 남기는 경고 로그를 줄인다.
logging.getLogger("hwp5").setLevel(logging.ERROR)

# HWPX 본문 텍스트 요소(<hp:t>)의 로컬 이름.
_HWPX_SECTION = re.compile(r"Contents/section\d+\.xml$", re.IGNORECASE)

# 임베드 텍스트가 이 길이 미만이면 스캔 PDF로 보고 OCR 로 폴백한다.
MIN_PDF_TEXT_CHARS = 50

# 서브셋 글꼴이 글리프를 유니코드 사설영역(PUA)에 매핑하고 ToUnicode 표를 안 넣으면,
# 추출기가 글자 대신 원본 글리프 코드(U+F639 등)를 돌려준다 — 숫자가 통째로 사라진다.
# 서로 다른 PUA 코드가 여러 개면 '본문이 그 글꼴로 그려졌다'는 신호다. 반대로 장식용
# 구분선처럼 같은 글리프만 반복되는 정상 문서는 distinct 가 1이라 걸리지 않는다
# (실제 볼트에서 확인: 레시피 노트 9종 = 손상, 견적서 구분선 1종 x98 = 정상).
MIN_DISTINCT_PUA = 5

# OCR 렌더링 해상도. 높을수록 정확하지만 느리고 메모리를 더 쓴다.
PDF_OCR_DPI = 300

# 스캔 문서 OCR 설정. 한국어 스캔본은 이진화(흑백 변환) 없이는 인식률이 매우 낮다.
# 임계값보다 어두운 픽셀만 검정으로 만들어 글자와 배경을 분리한다.
OCR_BINARIZE_THRESHOLD = 150
# psm 6 = "하나의 균일한 텍스트 블록으로 간주" — 본문 위주 계약서 스캔에 가장 정확했다.
PDF_OCR_CONFIG = "--psm 6"


def _pdf_embedded_text(file_path: Path) -> str:
    with pdfplumber.open(file_path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages).strip()


def _ocr_scanned_page(image: Image.Image) -> str:
    """스캔 페이지 이미지를 전처리(그레이스케일·대비·이진화)한 뒤 OCR 한다."""
    gray = ImageOps.autocontrast(image.convert("L"))
    binary = gray.point(lambda px: 0 if px < OCR_BINARIZE_THRESHOLD else 255, "1")
    return pytesseract.image_to_string(binary, lang=OCR_LANGUAGES, config=PDF_OCR_CONFIG)


_gemini = None


def _gemini_client():
    """google-genai 클라이언트를 최초 1회만 생성해 캐시한다(모듈 import 시엔 안 만든다)."""
    global _gemini
    if _gemini is None:
        from google import genai  # tesseract 전용 머신엔 없을 수 있어 지연 import

        api_key = os.getenv("GEMINI_API_KEY", "").strip()
        if not api_key:
            raise ValueError("GEMINI_API_KEY 가 비어 있습니다(.env 에 발급 키를 넣으세요).")
        _gemini = genai.Client(api_key=api_key)
    return _gemini


# 요약·번역·추측을 금지해야 한다. 모델이 '읽어서 정리'하기 시작하면 지번·금액 같은
# 원문 값이 조용히 바뀌어, 이 도메인에서 가장 위험한 종류의 오류가 된다.
GEMINI_OCR_PROMPT = (
    "이 이미지는 한국어 문서를 스캔한 것입니다. 보이는 모든 글자를 원문 그대로 옮겨 적으세요. "
    "표는 줄바꿈과 탭으로 구조를 유지하세요. 요약·설명·번역·추측을 하지 말고, 읽을 수 없는 "
    "글자는 생략하세요. 글자가 없으면 빈 문자열만 출력하세요."
)


def _gemini_ocr_page(image: Image.Image) -> str:
    """스캔 페이지 이미지를 Gemini 비전 모델로 전사한다(전처리 없이 원본 렌더 그대로).

    Tesseract 와 달리 이진화하지 않는다 — 비전 모델은 원본 계조에서 더 정확하다."""
    from google.genai import types

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    response = _gemini_client().models.generate_content(
        model=GEMINI_OCR_MODEL,
        contents=[
            types.Part.from_bytes(data=buffer.getvalue(), mime_type="image/png"),
            GEMINI_OCR_PROMPT,
        ],
        config=types.GenerateContentConfig(
            temperature=0.0,
            max_output_tokens=OCR_MAX_OUTPUT_TOKENS,
            thinking_config=types.ThinkingConfig(thinking_level="minimal"),
        ),
    )
    # 잘린 응답을 조용히 받아들이면 페이지 중간에서 끊긴 주소·금액이 '전부'로 둔갑한다.
    # 이 도메인에서 가장 위험한 실패라, 잘렸으면 예외로 올려 격리·알림에 걸리게 한다.
    candidate = (response.candidates or [None])[0]
    finish = getattr(candidate, "finish_reason", None)
    if getattr(finish, "name", str(finish)) == "MAX_TOKENS":
        raise ValueError(
            f"OCR 응답이 잘렸습니다(max_output_tokens={OCR_MAX_OUTPUT_TOKENS}). "
            "페이지가 너무 조밀하면 GEMINI_OCR_MAX_OUTPUT_TOKENS 를 올리세요."
        )
    return (response.text or "").strip()


def _ocr_page(image: Image.Image) -> str:
    """OCR_PROVIDER 에 따라 tesseract(로컬)/gemini(클라우드) 백엔드로 분기한다."""
    if OCR_PROVIDER == "gemini":
        return _gemini_ocr_page(image)
    return _ocr_scanned_page(image)


def _pdf_ocr_text(file_path: Path) -> str:
    """PDF 각 페이지를 이미지로 렌더링한 뒤 전처리·OCR 해 텍스트를 뽑는다."""
    pages = []
    with fitz.open(file_path) as doc:
        # gemini 백엔드는 페이지당 API 호출 1회라 상한을 둔다(로컬 tesseract 는 무제한).
        limit = GEMINI_OCR_MAX_PAGES if OCR_PROVIDER == "gemini" else len(doc)
        for index, page in enumerate(doc):
            if index >= limit:
                logger.warning(
                    "OCR 페이지 상한(%d) 초과, 앞 %d쪽까지만 읽음: %s (전체 %d쪽)",
                    limit, limit, file_path.name, len(doc),
                )
                break
            pix = page.get_pixmap(dpi=PDF_OCR_DPI)
            with Image.open(io.BytesIO(pix.tobytes("png"))) as image:
                pages.append(_ocr_page(image))
    return "\n".join(pages).strip()


def _has_unmapped_glyphs(text: str) -> bool:
    """사설영역(PUA) 문자가 여러 종류 섞여 있으면 임베드 텍스트를 믿을 수 없다는 뜻."""
    return len({c for c in text if unicodedata.category(c) == "Co"}) >= MIN_DISTINCT_PUA


def extract_pdf_text(file_path: Path) -> str:
    """임베드 텍스트를 우선 추출하고, 비어 있거나(스캔 PDF) 깨졌으면 OCR 로 폴백한다."""
    text = _pdf_embedded_text(file_path)
    garbled = _has_unmapped_glyphs(text)
    if garbled:
        # 페이지를 이미지로 읽으면 글꼴 인코딩과 무관하게 보이는 대로 숫자를 얻는다.
        logger.warning("임베드 텍스트가 글꼴 서브셋으로 깨져 OCR 로 다시 읽음: %s", file_path.name)
    elif len(text) >= MIN_PDF_TEXT_CHARS:
        return text

    # OCR 실패는 삼키지 않는다. 이 분기까지 왔다는 건 임베드 텍스트가 이미 50자 미만이라는
    # 뜻이라, 폴백해봐야 페이지번호·워터마크 몇 글자를 '문서 전체'로 분류하게 된다. 무인
    # 운영 머신에서 그건 조용한 오분류가 되므로, 예외를 그대로 올려 _failed 격리 + 알림에
    # 걸리게 한다(사람이 다시 넣으면 된다). OCR 이 '빈 결과'를 준 경우는 예외가 아니므로
    # 아래 비교에서 기존 임베드 텍스트가 그대로 살아남는다.
    ocr_text = _pdf_ocr_text(file_path)
    if garbled:
        # 깨진 임베드 텍스트는 길이만 길다. 더 짧아도 값이 살아 있는 OCR 결과가 낫다.
        return ocr_text or text
    return ocr_text if len(ocr_text) > len(text) else text


def extract_image_text(file_path: Path) -> str:
    with Image.open(file_path) as image:
        if OCR_PROVIDER == "gemini":
            return _gemini_ocr_page(image)
        return pytesseract.image_to_string(image, lang=OCR_LANGUAGES).strip()


def extract_plain_text(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8").strip()


def extract_hwp_text(file_path: Path) -> str:
    """한글 v5(.hwp) 바이너리에서 본문 텍스트를 추출한다(pyhwp)."""
    # 무거운 의존성이라 필요할 때만 임포트한다.
    from hwp5.hwp5txt import TextTransform
    from hwp5.xmlmodel import Hwp5File

    out = io.BytesIO()
    with closing(Hwp5File(str(file_path))) as hwp:
        TextTransform().transform_hwp5_to_text(hwp, out)
    return out.getvalue().decode("utf-8", "ignore").strip()


def extract_hwpx_text(file_path: Path) -> str:
    """한글 OWPML(.hwpx, zip+xml)에서 본문 텍스트를 추출한다."""
    parts = []
    with zipfile.ZipFile(file_path) as archive:
        sections = sorted(n for n in archive.namelist() if _HWPX_SECTION.search(n))
        for name in sections:
            root = ET.fromstring(archive.read(name))
            for element in root.iter():
                # 네임스페이스를 무시하고 로컬 이름이 't'(텍스트 런)인 요소만 모은다.
                if element.tag.rsplit("}", 1)[-1] == "t" and element.text:
                    parts.append(element.text)
    return "\n".join(parts).strip()


def extract_xlsx_text(file_path: Path) -> str:
    """엑셀 파일(.xlsx/.xls)에서 모든 시트의 셀 텍스트를 추출한다."""
    import openpyxl

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[시트: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append("\t".join(cells))
    wb.close()
    return "\n".join(parts).strip()


def extract_docx_text(file_path: Path) -> str:
    """워드 파일(.docx)에서 단락 텍스트를 추출한다."""
    import docx

    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 표(table) 안의 텍스트도 수집한다.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))
    return "\n".join(paragraphs).strip()


def _supported_attachment(name: str) -> bool:
    """첨부 파일명이 파이프라인이 처리할 수 있는 형식인지 판단한다.

    중첩 .msg(이메일 속 이메일)는 무한 재귀를 막기 위해 제외한다."""
    suffix = Path(name).suffix.lower()
    return suffix in SUPPORTED_EXTENSIONS and suffix not in MSG_EXTENSIONS


def _attachment_name(att) -> str | None:
    """첨부의 표시용 파일명(긴 이름 우선, 없으면 짧은 이름)."""
    return getattr(att, "longFilename", None) or getattr(att, "shortFilename", None)


def msg_attachments(file_path: Path) -> list[tuple[str, bytes]]:
    """.msg 이메일에서 '처리 가능한 첨부'를 (파일명, 바이트) 목록으로 돌려준다.

    파이프라인이 각 첨부를 독립 파일로 다시 처리해 자체 노트로 만들 수 있게 한다.
    smime.p7s·winmail.dat 같은 비문서, 중첩 이메일(.msg, data 가 bytes 아님),
    깨진/웹 첨부(data 가 None)는 건너뛴다."""
    import extract_msg

    out: list[tuple[str, bytes]] = []
    with extract_msg.openMsg(str(file_path)) as msg:
        for att in msg.attachments:
            data = getattr(att, "data", None)
            if not isinstance(data, (bytes, bytearray)):
                continue  # 중첩 .msg(MSGFile)·깨진/웹 첨부(None) 등은 제외
            name = _attachment_name(att)
            if name and _supported_attachment(name):
                out.append((name, bytes(data)))
    return out


def extract_msg_text(file_path: Path) -> str:
    """아웃룩 이메일(.msg)에서 제목·발신자·본문·첨부 목록을 추출한다."""
    import extract_msg

    with extract_msg.openMsg(str(file_path)) as msg:
        subject = msg.subject or ""
        sender = msg.sender or ""
        body = msg.body or ""
        attachments = [
            name for att in msg.attachments if (name := _attachment_name(att))
        ]
    lines = []
    if subject:
        lines.append(f"제목: {subject}")
    if sender:
        lines.append(f"발신자: {sender}")
    if body:
        lines.append(body.strip())
    if attachments:
        lines.append("첨부파일: " + ", ".join(attachments))
    return "\n".join(lines).strip()


# 전자세금계산서(NTS) XML 에서 뽑을 핵심 필드 — 로컬 태그 이름(네임스페이스 무시) → 라벨.
# 국세청 표준 스키마의 영문 태그명을 기준으로 하되, 없으면 일반 XML 폴백으로 넘어간다.
_TAX_INVOICE_FIELDS = {
    "IssueDate": "작성일자",
    "ChargeTotal": "공급가액",
    "TaxTotal": "세액",
    "GrandTotal": "합계금액",
}
# 공급자/공급받는자 같은 당사자 블록에서 뽑을 하위 필드.
_TAX_PARTY_FIELDS = {
    "NameName": "상호",
    "CompanyRegNumNumber": "사업자등록번호",
}


def extract_xml_text(file_path: Path) -> str:
    """구조화 XML(전자세금계산서 등)에서 텍스트를 추출한다.

    국세청 전자세금계산서 스키마를 알아보면 핵심 필드를 라벨과 함께 뽑아 분류기가
    category=세금계산서·상대방·날짜를 잡기 쉽게 한다. 인식 못 하는 스키마면 모든
    요소의 텍스트를 그대로 모으는 일반 XML 폴백으로 처리한다(extract_hwpx_text 와
    같은 네임스페이스-무시 방식)."""
    tree = ET.parse(file_path)
    root = tree.getroot()

    def _local(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    # 전자세금계산서 스키마 감지: 루트 또는 하위에 TaxInvoice 요소가 있으면 구조화 파싱.
    is_tax_invoice = any(_local(el.tag) == "TaxInvoice" for el in root.iter()) or (
        _local(root.tag) == "TaxInvoice"
    )

    if is_tax_invoice:
        lines = ["[전자세금계산서]"]
        # 당사자 블록(공급자/공급받는자)을 순서대로 라벨링한다.
        party_labels = ["공급자", "공급받는자"]
        parties = [el for el in root.iter() if _local(el.tag) == "Party"]
        for idx, party in enumerate(parties):
            label = party_labels[idx] if idx < len(party_labels) else f"당사자{idx + 1}"
            fields = []
            for el in party.iter():
                key = _TAX_PARTY_FIELDS.get(_local(el.tag))
                if key and el.text and el.text.strip():
                    fields.append(f"{key} {el.text.strip()}")
            if fields:
                lines.append(f"{label}: " + ", ".join(fields))
        # 금액·날짜 등 단일 필드.
        for el in root.iter():
            key = _TAX_INVOICE_FIELDS.get(_local(el.tag))
            if key and el.text and el.text.strip():
                lines.append(f"{key}: {el.text.strip()}")
        # 품목명(있으면).
        items = [
            el.text.strip()
            for el in root.iter()
            if _local(el.tag) == "ItemName" and el.text and el.text.strip()
        ]
        if items:
            lines.append("품목: " + ", ".join(items))
        text = "\n".join(lines).strip()
        # 라벨 헤더만 남고 실제 값이 하나도 없으면 일반 폴백으로 넘어간다.
        if len(lines) > 1:
            return text

    # 일반 XML 폴백: 네임스페이스를 무시하고 모든 요소의 텍스트를 모은다.
    parts = []
    for el in root.iter():
        if el.text and el.text.strip():
            parts.append(el.text.strip())
    return "\n".join(parts).strip()


def extract_text(file_path: Path) -> str:
    """파일 확장자에 따라 적절한 추출 방식을 선택해 텍스트를 반환한다."""
    suffix = file_path.suffix.lower()

    if suffix in PDF_EXTENSIONS:
        return extract_pdf_text(file_path)
    if suffix in IMAGE_EXTENSIONS:
        return extract_image_text(file_path)
    if suffix in TEXT_EXTENSIONS:
        return extract_plain_text(file_path)
    if suffix in HWP_EXTENSIONS:
        return extract_hwp_text(file_path)
    if suffix in HWPX_EXTENSIONS:
        return extract_hwpx_text(file_path)
    if suffix in XLSX_EXTENSIONS:
        return extract_xlsx_text(file_path)
    if suffix in DOCX_EXTENSIONS:
        return extract_docx_text(file_path)
    if suffix in MSG_EXTENSIONS:
        return extract_msg_text(file_path)
    if suffix in XML_EXTENSIONS:
        return extract_xml_text(file_path)

    raise ValueError(f"지원하지 않는 파일 형식입니다: {suffix}")
