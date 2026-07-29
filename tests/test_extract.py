import zipfile
from pathlib import Path

import openpyxl
import pytest
import docx

from extractors.extract import extract_text

FIXTURES = Path(__file__).parent / "fixtures"


def _make_hwpx(path: Path, paragraphs: list[str]) -> None:
    """테스트용 최소 HWPX(zip+OWPML) 파일을 만든다."""
    runs = "".join(f'<hp:t xmlns:hp="x">{p}</hp:t>' for p in paragraphs)
    section = f'<?xml version="1.0"?><hp:sec xmlns:hp="x">{runs}</hp:sec>'
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("Contents/section0.xml", section)


def test_extract_plain_text():
    assert extract_text(FIXTURES / "sample.txt") == "안녕하세요, 테스트 텍스트입니다."


def test_extract_pdf_text(monkeypatch):
    """임베드 텍스트가 짧아 OCR 로 넘어가도, OCR 이 빈 결과면 임베드 텍스트가 살아남는다.

    OCR 백엔드는 머신마다 다르고(gemini 는 네트워크·요금을 탄다) 테스트가 그것에 의존하면
    안 되므로 여기서 격리한다."""
    from extractors import extract

    monkeypatch.setattr(extract, "_pdf_ocr_text", lambda _: "")
    assert "Hello PDF" in extract_text(FIXTURES / "sample.pdf")


def test_font_subset_garbage_falls_back_to_ocr(monkeypatch):
    """숫자가 사설영역(PUA) 글리프로 나오는 PDF는 임베드 텍스트를 믿지 않고 OCR 로 다시 읽는다.

    실제 사례: 레시피 PDF 의 영양정보·조리순서 숫자가 U+F639~U+F640 으로 추출됐다."""
    from extractors import extract

    garbled = (
        "Calories \nFat g\nCarbs g\n"
        "Fiber g\nSugar g\nProtein g\n" + "설명 " * 40
    )
    monkeypatch.setattr(extract, "_pdf_embedded_text", lambda _: garbled)
    monkeypatch.setattr(extract, "_pdf_ocr_text", lambda _: "Calories 711")
    # 깨진 텍스트가 더 길어도 OCR 결과를 택해야 한다.
    assert extract_text(FIXTURES / "sample.pdf") == "Calories 711"


def test_decorative_pua_glyph_does_not_trigger_ocr(monkeypatch):
    """같은 PUA 글리프만 반복되는 장식용 구분선은 정상 문서다 — OCR 로 넘기지 않는다.

    실제 사례: 견적서 노트의 U+F000 x98 은 가로줄 장식이고 숫자는 멀쩡했다."""
    from extractors import extract

    text = "견적서\n" + "" * 98 + "\n합계 10,000,000원\n" + "항목 " * 40
    monkeypatch.setattr(extract, "_pdf_embedded_text", lambda _: text)
    monkeypatch.setattr(
        extract, "_pdf_ocr_text", lambda _: pytest.fail("정상 문서를 OCR 로 넘기면 안 된다")
    )
    assert "10,000,000원" in extract_text(FIXTURES / "sample.pdf")


def test_extract_pdf_propagates_ocr_failure(monkeypatch):
    """OCR 실패를 삼키면 워터마크 몇 글자가 '문서 전체'가 된다 — 예외를 그대로 올려야 한다.

    (호출부인 pipeline.process_file 이 이 예외를 받아 _failed 격리 + 알림까지 처리한다.)"""
    from extractors import extract

    def _boom(_):
        raise RuntimeError("OCR 백엔드 실패")

    monkeypatch.setattr(extract, "_pdf_ocr_text", _boom)
    with pytest.raises(RuntimeError):
        extract_text(FIXTURES / "sample.pdf")


def _fake_gemini_client(finish_name: str, text: str):
    """finish_reason 과 본문을 지정한 가짜 google-genai 클라이언트를 만든다."""
    finish = type("Finish", (), {"name": finish_name})()
    candidate = type("Candidate", (), {"finish_reason": finish})()
    response = type("Response", (), {"candidates": [candidate], "text": text})()
    models = type("Models", (), {"generate_content": lambda self, **kw: response})()
    return type("Client", (), {"models": models})()


def test_gemini_ocr_rejects_truncated_response(monkeypatch):
    """응답이 잘렸는데(MAX_TOKENS) 조용히 받으면 끊긴 주소·금액이 전부로 둔갑한다."""
    from PIL import Image

    from extractors import extract

    monkeypatch.setattr(
        extract,
        "_gemini_client",
        lambda: _fake_gemini_client("MAX_TOKENS", "서울특별시 성동구 성수동1가 685-"),
    )
    with pytest.raises(ValueError, match="잘렸"):
        extract._gemini_ocr_page(Image.new("RGB", (4, 4)))


def test_gemini_ocr_returns_complete_response(monkeypatch):
    """정상 종료(STOP)면 전사 결과를 그대로 돌려준다(잘림 검사가 오탐하지 않는다)."""
    from PIL import Image

    from extractors import extract

    monkeypatch.setattr(
        extract, "_gemini_client", lambda: _fake_gemini_client("STOP", " 685-317 대수선 \n")
    )
    assert extract._gemini_ocr_page(Image.new("RGB", (4, 4))) == "685-317 대수선"


def test_ocr_page_dispatches_to_gemini(monkeypatch):
    """OCR_PROVIDER=gemini 면 Tesseract 대신 비전 백엔드로 분기한다."""
    from PIL import Image

    from extractors import extract

    monkeypatch.setattr(extract, "OCR_PROVIDER", "gemini")
    monkeypatch.setattr(extract, "_gemini_ocr_page", lambda image: "성수동1가 685-317")
    assert extract._ocr_page(Image.new("RGB", (4, 4))) == "성수동1가 685-317"


def test_gemini_ocr_stops_at_page_cap(tmp_path, monkeypatch):
    """페이지당 API 1회이므로, 상한을 넘는 스캔은 앞쪽까지만 읽는다."""
    import fitz

    from extractors import extract

    pdf = tmp_path / "long_scan.pdf"
    doc = fitz.open()
    for _ in range(5):
        doc.new_page()
    doc.save(str(pdf))
    doc.close()

    calls = []
    monkeypatch.setattr(extract, "OCR_PROVIDER", "gemini")
    monkeypatch.setattr(extract, "GEMINI_OCR_MAX_PAGES", 2)
    monkeypatch.setattr(
        extract, "_gemini_ocr_page", lambda image: calls.append(1) or "쪽"
    )
    extract._pdf_ocr_text(pdf)
    assert len(calls) == 2


def test_extract_hwpx_text(tmp_path):
    hwpx = tmp_path / "sample.hwpx"
    _make_hwpx(hwpx, ["계약서 제목", "제1조 목적"])
    text = extract_text(hwpx)
    assert "계약서 제목" in text
    assert "제1조 목적" in text


def test_extract_xlsx_text(tmp_path):
    xlsx = tmp_path / "sample.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "견적"
    ws.append(["항목", "금액"])
    ws.append(["철근공사", "5000000"])
    wb.save(xlsx)
    text = extract_text(xlsx)
    assert "견적" in text
    assert "철근공사" in text
    assert "5000000" in text


def test_extract_docx_text(tmp_path):
    docx_file = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("공사 계약서")
    doc.add_paragraph("제1조 총칙")
    doc.save(str(docx_file))
    text = extract_text(docx_file)
    assert "공사 계약서" in text
    assert "제1조 총칙" in text


def test_extract_xml_tax_invoice(tmp_path):
    """국세청 전자세금계산서 스키마(네임스페이스 포함)에서 핵심 필드를 뽑는다."""
    xml = tmp_path / "tax.xml"
    xml.write_text(
        """<?xml version="1.0" encoding="UTF-8"?>
<TaxInvoice xmlns="urn:kr:or:nts:standard">
  <ExchangedDocument><IssueDate>2026-02-04</IssueDate></ExchangedDocument>
  <TaxInvoiceTradeSettlement>
    <SellerParty><Party>
      <SpecifiedOrganization><NameName>누리구조</NameName></SpecifiedOrganization>
      <CompanyRegNumNumber>1234567890</CompanyRegNumNumber>
    </Party></SellerParty>
    <BuyerParty><Party>
      <SpecifiedOrganization><NameName>케이씨건설</NameName></SpecifiedOrganization>
      <CompanyRegNumNumber>9876543210</CompanyRegNumNumber>
    </Party></BuyerParty>
    <ChargeTotal>5000000</ChargeTotal>
    <TaxTotal>500000</TaxTotal>
    <GrandTotal>5500000</GrandTotal>
  </TaxInvoiceTradeSettlement>
  <TaxInvoiceTradeLineItem><ItemName>구조설계 계약금</ItemName></TaxInvoiceTradeLineItem>
</TaxInvoice>
""",
        encoding="utf-8",
    )
    text = extract_text(xml)
    assert "전자세금계산서" in text
    assert "누리구조" in text
    assert "케이씨건설" in text
    assert "5500000" in text
    assert "2026-02-04" in text
    assert "구조설계 계약금" in text


def test_extract_xml_generic_fallback(tmp_path):
    """세금계산서가 아닌 일반 XML 은 모든 요소 텍스트를 모은다."""
    xml = tmp_path / "note.xml"
    xml.write_text(
        '<?xml version="1.0"?><root><memo>회의 메모</memo><item>안건 검토</item></root>',
        encoding="utf-8",
    )
    text = extract_text(xml)
    assert "회의 메모" in text
    assert "안건 검토" in text


def test_supported_attachment_filters_types():
    from extractors.extract import _supported_attachment

    assert _supported_attachment("철거해체계약서.pdf")
    assert _supported_attachment("내역서.xlsx")
    assert _supported_attachment("공문.hwp")
    # 비문서·중첩 이메일은 제외(무한 재귀·노이즈 방지).
    assert not _supported_attachment("smime.p7s")
    assert not _supported_attachment("winmail.dat")
    assert not _supported_attachment("forwarded.msg")


def test_extract_unsupported_extension():
    with pytest.raises(ValueError):
        extract_text(Path("unsupported.xyz"))
